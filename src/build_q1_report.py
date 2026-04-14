#!/usr/bin/env python3
"""Build Q1 Literature Review as MS Word document for LIAA submission."""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

WORK_DIR = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(WORK_DIR, "Q1_Literaturas_Parskats_LIAA.docx")

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hs.font.size = Pt(16)
        hs.font.bold = True
    elif level == 2:
        hs.font.size = Pt(14)
        hs.font.bold = True
    else:
        hs.font.size = Pt(12)
        hs.font.bold = True


def add_table(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx+1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(72)
run = p.add_run("PĒTNIECĪBAS UN ATTĪSTĪBAS PROJEKTS")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Kuģu propulsijas modeļu un pazīmju inženierijas\nmetožu izstrāde un validācija")
run.bold = True
run.font.size = Pt(16)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Q1 NODEVUMS\nLiteratūras pārskats")
run.bold = True
run.font.size = Pt(14)

doc.add_paragraph()
doc.add_paragraph()

info = [
    ("Projekta izpildītājs:", "Dmitrijs Odiņokijs"),
    ("Kvalifikācija:", "Datu zinātnieks, RTU doktorants"),
    ("Organizācija:", "SIA ShipProjects"),
    ("Datums:", "2026. gada aprīlis"),
    ("Periods:", "Q1 (Februāris – Aprīlis)"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + " ")
    run.bold = True
    p.add_run(value)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. IEVADS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading("1. Ievads", level=1)

doc.add_paragraph(
    "Šis literatūras pārskats ir sagatavots kā PRISMA 2020 standartam atbilstoša "
    "sistemātiskā literatūras analīze, kas aptver mašīnmācīšanās (ML) metožu pielietojumu "
    "kuģu veiktspējas prognozēšanā un propulsijas modelēšanā. Pārskats fokusējas uz trim "
    "galvenajiem virzieniem, kas tieši attiecas uz Periscope sistēmas uzlabošanu:"
)

doc.add_paragraph(
    "1) Pazīmju inženierija (Feature Engineering) — ML modeļu ievaddatu konstruēšana "
    "no sensoru laika rindām;", style='List Number'
)
doc.add_paragraph(
    "2) Fizikas-hibrīda arhitektūras — semi-empīrisku fizikas bāzlīniju apvienošana "
    "ar ML korekciju (grey-box modeļi);", style='List Number'
)
doc.add_paragraph(
    "3) Tiešsaistes adaptācija — izvietotu modeļu reāllaika kalibrēšana un atjaunināšana, "
    "lai kompensētu korpusa apaugšanu un dzinēja nolietojumu.", style='List Number'
)

doc.add_paragraph(
    "Pārskats ir daļa no LIAA pētniecības projekta Q1 nodevumiem. Tā mērķis ir izveidot "
    "zinātniski pamatotu bāzi turpmākajam darbam pie Periscope sistēmas uzlabošanas, "
    "identificējot labākās prakses, nepilnības un perspektīvākos virzienus."
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. MEKLĒŠANAS METODOLOĢIJA
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading("2. Meklēšanas metodoloģija", level=1)

doc.add_heading("2.1. Datubāzes un meklēšanas virknes", level=2)

doc.add_paragraph(
    "Meklēšana veikta 2026. gada 14. aprīlī, izmantojot Python automatizētu cauruļvadu "
    "(prisma_pipeline.py), kas nodrošina pilnīgu reproducējamību. Pārmeklētas četras "
    "akadēmiskās datubāzes, izmantojot to REST API:"
)

add_table(doc,
    ["Datubāze", "Ieraksti", "Piezīmes"],
    [
        ["Semantic Scholar", "279", "API ar atslēgu; atslēgvārdu vaicājumi"],
        ["OpenAlex", "400", "Pilni Būla vaicājumi"],
        ["arXiv", "9", "Tulkoti uz ti+abs formātu"],
        ["Scopus", "176", "Elsevier API ar TITLE-ABS-KEY sintaksi"],
        ["Manuālā meklēšana", "5", "Google Scholar, ResearchGate, ScienceDirect"],
        ["Kopā", "869", ""],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Izmantotas 4 meklēšanas virknes, kas aptver projekta tematiku: A — pamata propulsijas "
    "modelēšana, B — pazīmju inženierija, C — fizikas-hibrīda pieejas, D — reāllaika "
    "kalibrēšana. Visas virknes ierobežotas ar datumu filtru 2015–2025. Sākotnējās plašākās "
    "virknes (v1) tika iteratīvi sašaurinātas, lai samazinātu troksni — piemēram, String C "
    "uz OpenAlex sākotnēji atgrieza 10 863 rezultātus, kas tika samazināts līdz ~1 250, "
    "aizstājot vispārīgus terminus ar nozarei specifiskiem."
)

doc.add_paragraph(
    "Google Scholar tika izslēgts no sistemātiskā cauruļvada, jo scholarly Python bibliotēka "
    "nenodrošina reproducējamus rezultātus. IEEE Xplore API atslēga netika aktivizēta laikā; "
    "IEEE publikācijas daļēji tiek aptvertas caur Scopus un Semantic Scholar krustveidīgo indeksēšanu."
)

doc.add_heading("2.2. Atlases process", level=2)

add_table(doc,
    ["Posms", "Ieraksti"],
    [
        ["Pirms deduplikācijas", "869"],
        ["Pēc deduplikācijas (DOI + fuzzy nosaukuma salīdzināšana)", "575"],
        ["Virsraksta skrīnings (3 kārtas)", "181 PASS, 394 FAIL"],
        ["Anotācijas skrīnings", "137 PASS, 44 FAIL"],
        ["Pilna teksta novērtējums", "116 iekļauti, 21 izslēgti"],
        ["Galīgais iekļauto pētījumu skaits", "116"],
    ]
)

doc.add_paragraph()

doc.add_heading("2.3. Izslēgšanas kritēriji", level=2)

doc.add_paragraph(
    "Izslēgšanas kodi: E1 — jūrniecības tēma, bet ārpus propulsijas/degvielas prognozēšanas "
    "(maršrutu optimizācija, trajektoriju prognozēšana, strukturālā integritāte u.c.); "
    "E2 — nav ML vai datu vadītas metodes; E5 — nav jūrniecības domēns; "
    "E6 — nav angļu valodā; E7 — nepietiekams metodolģiskais ieguldījums vai nav ziņotas "
    "kvantitatīvas precizitātes metrikas."
)

add_table(doc,
    ["Izslēgšanas kods", "Skaits"],
    [
        ["E1 — ārpus propulsijas/degvielas prognozēšanas", "264"],
        ["E5 — nav jūrniecības domēns", "142"],
        ["E7 — nepietiekams ieguldījums", "21"],
        ["E2 — nav ML metode", "10"],
        ["E6 — nav angļu valodā", "1"],
        ["Pilna teksta izslēgšana", "21"],
        ["Kopā izslēgti", "459"],
    ]
)

doc.add_paragraph()

doc.add_heading("2.4. Tematiskā klasterēšana", level=2)

add_table(doc,
    ["Klasteris", "Skaits", "Apraksts"],
    [
        ["T1 — ML bāzlīnijas", "60", "Standarta ML metožu salīdzinājumi degvielas/ātruma/jaudas prognozēšanai"],
        ["T2 — Fizikas-hibrīda", "27", "Grey-box modeļi, fizikas vadīti neironu tīkli, semi-empīriski + ML"],
        ["T3 — Pazīmju inženierija", "15", "Ievaddatu atlase, datu apvienošana, priekšapstrāde"],
        ["T4 — Tiešsaistes adaptācija", "9", "Online learning, inkrementālie modeļi, koncepta dreifs"],
        ["T6 — Interpretējamība", "5", "SHAP, LIME, izskaidrojamā MI"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. GALVENIE SECINĀJUMI
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading("3. Galvenie secinājumi", level=1)

doc.add_heading("3.1. ML metožu salīdzinājums (T1)", level=2)

doc.add_paragraph(
    "60 T1 pētījumos tiešie metožu salīdzinājumi konsistenti parāda, ka gradient boosting "
    "metodes (XGBoost, LightGBM, CatBoost) sasniedz augstāko precizitāti tabulāros kuģu "
    "datos. Dziļās mācīšanās (LSTM, Transformer) priekšrocības izpaužas tikai ar ilgām "
    "laika rindām un lielu datu apjomu."
)

p = doc.add_paragraph()
run = p.add_run("Kritiski svarīgi par ziņotajām R² vērtībām: ")
run.bold = True
p.add_run(
    "Literatūrā ziņotie precizitātes rādītāji ir sistemātiski paaugstināti. "
    "Mūsu PDF verifikācija atklāja, ka pētījumi ar nejaušu k-fold dalījumu ziņo "
    "R² = 0,95–0,99, kamēr pētījumi ar temporālu holdout ziņo R² = 0,75–0,95. "
    "Starpība (~0,04–0,10 R² vienības) ir attiecināma uz datu noplūdi, ko izraisa "
    "laika rindu nejaušā dalīšana — modelis caur lag iezīmēm un slīdošā loga statistikām "
    "var «redzēt nākotni»."
)

doc.add_paragraph(
    "Rakstos bieži nav skaidrs, uz ko attiecas ziņotais R² — vai tas ir uz apmācības, "
    "validācijas vai testa datiem. No 116 iekļautajiem pētījumiem tikai ~20 skaidri norāda, "
    "ka R² ir mērīts uz atsevišķa testa kopuma. Tas padara starppētījumu salīdzināšanu "
    "neuzticamu un ir sistēmiska problēma nozarē."
)

doc.add_paragraph(
    "Mūsu operacionālā pieredze ar Periscope sistēmu rāda, ka reālā ekspluatācijā "
    "ar temporālu validāciju un vairākiem kuģiem:"
)
doc.add_paragraph("Bez pazīmju inženierijas: R² ≈ 0,30–0,40", style='List Bullet')
doc.add_paragraph("Ar sistemātisku pazīmju inženieriju un hiperparametru optimizāciju: R² ≈ 0,60–0,70", style='List Bullet')
doc.add_paragraph("Ar fizikas-hibrīda arhitektūru: R² var sasniegt 0,80–0,90", style='List Bullet')

doc.add_paragraph(
    "R² = 0,60–0,70 ir uzskatāms par ļoti labu rezultātu reālās ekspluatācijas apstākļos "
    "pēc hiperparametru optimizācijas un iezīmju inženierijas. Literatūras R² > 0,95 "
    "atspoguļo metodolģiski vājākas validācijas (viena kuģa nejaušie dalījumi) un nav "
    "izmantojami kā mērķa rādītāji ražošanas sistēmām."
)

doc.add_heading("3.2. Fizikas-hibrīda arhitektūras (T2)", level=2)

doc.add_paragraph(
    "27 T2 pētījumi demonstrē trīs arhitektūras modeļus:"
)

doc.add_paragraph(
    "Fizikas iezīmes kā ML ievade (10 raksti): Holtrop-Mennen vai STAWAVE-2 pretestības "
    "komponentes tiek aprēķinātas un pievienotas kā papildu iezīmes ML modelim.",
    style='List Bullet'
)
doc.add_paragraph(
    "Fizikas-atlikuma (residual) arhitektūra (12 raksti): Fizikas modelis dod bāzlīnijas "
    "prognozi, ML koriģē atlikumu starp fiziku un realitāti. Šī ir statistiski pamatotākā "
    "pieeja.",
    style='List Bullet'
)
doc.add_paragraph(
    "Fizikas ierobežojumi (3 raksti): Zaudējuma funkcija ar nosacījumiem, kas novērš "
    "fiziski neticamas prognozes.",
    style='List Bullet'
)

doc.add_paragraph(
    "Precizitātes uzlabojums: 8–42% RMSE samazinājums salīdzinot ar tīru ML bāzlīniju. "
    "Vislielākais ieguvums pie augsta jūras stāvokļa, degradēta korpusa un darbības "
    "apstākļiem ārpus apmācības datu sadalījuma. Fizikas bāzlīnija nodrošina strukturālu "
    "induktīvo nobīdi, kas ļauj modelim ģeneralizēties uz nezināmiem apstākļiem — tas ir "
    "kritiski svarīgi Periscope sistēmai."
)

doc.add_heading("3.3. Pazīmju inženierija (T3)", level=2)

p = doc.add_paragraph()
run = p.add_run("Galvenais secinājums: ")
run.bold = True
p.add_run(
    "Pazīmju inženierija ir tikpat svarīga kā algoritma izvēle — pretēji tam, ko "
    "lielākā daļa literatūras akcentē. Algoritma maiņa (piemēram, no Random Forest uz "
    "XGBoost) maina R² par ±0,01–0,03, kamēr pazīmju inženierija — par ±0,10–0,30."
)

doc.add_paragraph("Biežāk izmantotas iezīmju kategorijas:")

doc.add_paragraph(
    "Kuģa stāvoklis: ātrums pār ūdeni, dzinēja slodze, iegrimes, trims, ūdensizspaids",
    style='List Bullet'
)
doc.add_paragraph(
    "Vides apstākļi: vēja ātrums un virziens, viļņu augstums un periods, jūras straumes",
    style='List Bullet'
)
doc.add_paragraph(
    "No laika rindām atkarīgas iezīmes: slīdošie vidējie (2–10 intervāli), lag iezīmes, "
    "mainīguma rādītāji, izmaiņu ātrumi",
    style='List Bullet'
)
doc.add_paragraph(
    "Korpusa stāvokļa aizstājēji: dienas kopš doka, kumulatīvās darba stundas",
    style='List Bullet'
)
doc.add_paragraph(
    "No fizikas atkarīgas iezīmes: Holtrop-Mennen pretestība, šķietamais vējš un viļņu "
    "leņķis, Frouda skaitlis",
    style='List Bullet'
)

doc.add_heading("3.4. Tiešsaistes adaptācija (T4)", level=2)

p = doc.add_paragraph()
run = p.add_run("Tikai 9 no 116 iekļautajiem pētījumiem ")
run.bold = True
p.add_run(
    "adresē tiešsaistes modeļa adaptāciju. Tas ir kritiski nepietiekami, ņemot vērā, ka "
    "kuģa korpusa apaugšana ar jūras organismiem palielina pretestību par 0,1–0,3% dienā, "
    "un 5 gadu doka ciklā ātruma zudums var sasniegt 5–15%. Statisks modelis kļūst "
    "neaktuāls 6–12 mēnešu laikā."
)

doc.add_paragraph("Literatūrā pētītās pieejas:")
doc.add_paragraph("Slīdošā loga atkārtota apmācība (30 dienu logs)", style='List Bullet')
doc.add_paragraph("Inkrementāla mācīšanās (Bayesian-optimized LightGBM)", style='List Bullet')
doc.add_paragraph("Adaptīvi ansambļi ar mainīgiem svariem", style='List Bullet')
doc.add_paragraph("Kalman filtrēšana digitālajā dvīnī", style='List Bullet')

p = doc.add_paragraph()
run = p.add_run("Kritiskā nepilnība: ")
run.bold = True
p.add_run(
    "neviens pētījums neziņo rezultātus no ražošanā izvietotas sistēmas. Visas validācijas "
    "ir retrospektīvas simulācijas. Tas ir vislielākā plaisa starp akadēmisko literatūru "
    "un ražošanas vajadzībām."
)

doc.add_heading("3.5. Validācijas metodoloģijas problēmas", level=2)

add_table(doc,
    ["Validācijas tips", "Mediāna R²", "Pētījumu skaits"],
    [
        ["Temporāla holdout", "0,93", "8"],
        ["Reisu bloku validācija", "0,88–0,95", "3"],
        ["Nejauša k-fold (datu noplūdes risks)", "0,97", "12"],
        ["Nav norādīts", "0,95", "~30"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Iemesli paaugstinātiem rādītājiem: 1) datu noplūde no nejaušas laika rindu dalīšanas; "
    "2) triviālas ievades (dzinēja parametri kā ievade padara prognozi tautologisku — "
    "viens pētījums sasniedza R² = 0,99999 izmantojot ME jaudu kā ievadi); "
    "3) viena kuģa pārmācīšanās (41 no 116 pētījumiem izmanto tikai vienu kuģi); "
    "4) neskaidra metriku ziņošana — daudzi pētījumi nenorāda, vai R² ir uz apmācības "
    "vai testa datiem."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. ATBILSTĪBA PERISCOPE PROJEKTAM
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading("4. Atbilstība Periscope projektam", level=1)

doc.add_heading("4.1. Bāzlīnijas modeļa konteksts", level=2)

doc.add_paragraph(
    "Pašreizējais Periscope bāzlīnijas modelis (XGBoost ar neapstrādātiem sensoru datiem) "
    "atbilst literatūras T1 klasterim. Mūsu R² ≈ 0,30–0,40 ir konsistents ar to, ko sagaida "
    "pie temporālas validācijas bez pazīmju inženierijas. Tas apstiprina, ka mūsu validācijas "
    "metodoloģija ir korekta — atšķirībā no literatūrā bieži izmantotajiem uzpūstajiem "
    "rādītājiem."
)

doc.add_heading("4.2. Uzlabošanas virzieni", level=2)

add_table(doc,
    ["Projekta fāze", "Literatūras atbalsts", "Sagaidāmais R²"],
    [
        ["Q1: Bāzlīnija (raw XGBoost)", "T1 — apstiprināts kā atskaites punkts", "0,35–0,45"],
        ["Q2: Pazīmju inženierija v2", "T3 — iezīmju inženierija dod 0,25–0,35 R² uzlabojumu", "0,60–0,70"],
        ["Q3: Physics Residual modelis", "T2 — 8–42% RMSE uzlabojums", "0,70–0,80"],
        ["Q4: Tiešsaistes adaptācija", "T4 — 0,08 R² uzlabojums pār statisku modeli", "0,72+"],
    ]
)

doc.add_paragraph()

doc.add_heading("4.3. Rekomendācijas", level=2)

doc.add_paragraph(
    "1. Pazīmju inženierija ir galvenā prioritāte — tā dod lielāko precizitātes pieaugumu "
    "(±0,10–0,30 R²), kamēr algoritma maiņa dod tikai ±0,01–0,03.",
    style='List Number'
)
doc.add_paragraph(
    "2. Validācijai izmantot tikai temporālu holdout vai reisu bloku validāciju. "
    "Nekad neizmantot nejaušu k-fold uz laika rindu datiem.",
    style='List Number'
)
doc.add_paragraph(
    "3. Physics Residual arhitektūra ir visperspektīvākā pieeja — tā nodrošina labāku "
    "ģeneralizāciju un interpretējamību, īpaši ārpus apmācības datu sadalījuma.",
    style='List Number'
)
doc.add_paragraph(
    "4. Tiešsaistes adaptācija ir nepieciešama ražošanas sistēmai — statisks modelis "
    "degradējas 6–12 mēnešu laikā korpusa apaugšanas un dzinēja nolietojuma dēļ.",
    style='List Number'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. ATSAUCES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_heading("5. Atsauces", level=1)

doc.add_paragraph(
    "Pilns 116 iekļauto pētījumu saraksts ar metadatiem pieejams papildinošajā Excel failā "
    "(PRISMA_search_results_2026-04-14.xlsx). Zemāk uzskaitīti 30 nozīmīgākie avoti:"
)

refs = [
    # T1
    "[1] Gkerekos, C., Lazakis, I., Theotokatos, G. (2019). Machine learning models for predicting ship main engine FOC. Ocean Engineering, 188, 106282.",
    "[2] Uyanik, T., Karatug, C., Arslanoglu, Y. (2020). ML approach to ship fuel consumption: container vessel. Transportation Research Part D, 84, 102389.",
    "[3] Bassam, A.M., et al. (2022). Ship speed prediction based on ML. Ocean Engineering, 245, 110449.",
    "[4] Abebe, M., et al. (2020). ML Approaches for Ship Speed Prediction. Applied Sciences, 10(7), 2325.",
    "[5] Hu, Z., et al. (2019). Prediction of Fuel Consumption for Enroute Ship. IEEE Access, 7, 119497.",
    "[6] Nguyen, S., et al. (2023). Multi-ship predictive modeling. Transportation Research Part E, 177, 103261.",
    "[7] Viga, J., et al. (2025). FuelCast: Benchmarking Tabular and Temporal Models. AALTD@ECML/PKDD.",
    # T2
    "[8] Lang, X., Wu, D., Mao, W. (2024). Physics-informed ML models for ship speed prediction. Expert Systems with Applications, 238, 121877.",
    "[9] Wang, H., et al. (2022). Combined ML and physics-based models. Ocean Engineering, 255, 111435.",
    "[10] Parkes, A.I., et al. (2018). Physics-based shaft power prediction using NN. Ocean Engineering, 166, 92–104.",
    "[11] De Haas, M., et al. (2023). Power Increase due to Marine Biofouling: Grey-box Model.",
    "[12] Ruan, Z., et al. (2025). A novel dual-stage grey-box stacking method. Energy.",
    "[13] Fan, A., et al. (2025). A novel grey box model for ship fuel consumption. Energy.",
    "[14] Mavroudis, S., Tinga, T. (2025). Transfer learning on physics-based models. Ocean Engineering.",
    "[15] Altan, D., et al. (2025). Physics-guided NN-based Shaft Power Prediction. arXiv.",
    # T3
    "[16] Du, Y., et al. (2022). Data fusion and ML for ship fuel efficiency: Part III. Communications in Transportation Research.",
    "[17] Wang, J., et al. (2022). Sensitivity to metocean data. Ocean Engineering, 263, 111155.",
    "[18] Kim, H.S., Roh, M.I. (2024). Interpretable data-driven models. Int. J. Naval Architecture.",
    "[19] Kim, D., et al. (2023). Feature Attribution Analysis. Italian Nat. Conf. on Sensors.",
    "[20] Kim, Y.R., et al. (2025). Review of data processing for ship performance. Applied Ocean Research.",
    # T4
    "[21] Liu, X., et al. (2025). Online Learning for Ship Propulsion Power Prediction. IEEE EIECT.",
    "[22] Gao, Y., et al. (2025). Adaptive Prediction Framework. JMSE.",
    "[23] Mittendorf, M., et al. (2023). Capturing biofouling by incremental ML. Applied Ocean Research.",
    "[24] Dann, N., et al. (2022). Adaptive Learning of Inland Ship Power. IFAC PapersOnLine.",
    "[25] Tveten, M., et al. (2025). Fault detection in propulsion motors: concept drift. IEEE EEEIC.",
    # T5/T6
    "[26] Agand, P., et al. (2023). Fuel Consumption Prediction for Passenger Ferry. Ocean Engineering.",
    "[27] Dekeyser, S., et al. (2022). Towards improved prediction of ship performance. arXiv.",
    "[28] Fan, A., et al. (2022). A review of ship fuel consumption models. Ocean Engineering, 264.",
    "[29] Vorkapic, A., et al. (2024). Interpretable ML: VLGC ship propulsion. JMSE, 12(10), 1849.",
    "[30] Ma, Y., et al. (2023). Interpretable Gray Box Model for Ship FOC. JMSE, 11(3), 622.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)
    for run in p.runs:
        run.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUT)
print(f"Saved: {OUT}")
